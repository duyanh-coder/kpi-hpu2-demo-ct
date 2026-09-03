#!/usr/bin/env python3
"""Parse KHCT data from QD 1480 docx and generate khct-catalog.json."""
import json, re, sys, os, time, random
sys.stdout.reconfigure(encoding='utf-8')

import docx

DOC_PATH = r'D:\kpi-hpu2-demo-ct\Docs\QD ban hanh bo KPI chinh, KHCT, KHDT 2026-2027-final.signed.signed.signed.signed.docx'
OUT_PATH = r'D:\kpi-hpu2-demo-ct\src\data\khct-catalog.json'

MONTHS = [
    ("8/2026", "Tháng 8 - 2026"),
    ("9/2026", "THÁNG 9 - 2026"),
    ("10/2026", "THÁNG 10 - 2026"),
    ("11/2026", "THÁNG 11 - 2026"),
    ("12/2026", "THÁNG 12 - 2026"),
    ("1/2027", "THÁNG 01 - 2027"),
    ("2/2027", "THÁNG 02 - 2027"),
    ("3/2027", "THÁNG 3 - 2027"),
    ("4/2027", "THÁNG 4 - 2027"),
    ("5/2027", "THÁNG 5 - 2027"),
    ("6/2027", "THÁNG 6 - 2027"),
    ("7/2027", "THÁNG 7 - 2027"),
]

def clean(s):
    return re.sub(r'\s+', ' ', s).strip()

def extract_khct():
    doc = docx.Document(DOC_PATH)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    tables = doc.tables

    # Identify KHCT tables: 8 cols, has task numbering in col 1
    khct_tables = []
    for ti, table in enumerate(tables):
        if len(table.columns) != 8:
            continue
        for row in table.rows[1:3] if len(table.rows) > 1 else []:
            cells = [clean(c.text) for c in row.cells]
            if re.match(r'^\d+\.$', cells[1] if len(cells) > 1 else ''):
                khct_tables.append(ti)
                break

    print(f"Total tables: {len(tables)}, KHCT tables: {len(khct_tables)}")

    # Column mapping: 0=field, 1=order, 2=taskName, 3=responsibleUnit,
    #                 4=coordinatingUnits, 5=kpiCodes, 6=deliverable, 7=deadline
    all_tasks = []
    current_field = ""

    for ti in khct_tables:
        table = tables[ti]
        for ri, row in enumerate(table.rows):
            cells = [clean(c.text) for c in row.cells]

            # Skip header rows
            if cells[0] == 'Lĩnh vực công tác' or cells[1] == 'Tên nhiệm vụ':
                continue

            order_match = re.match(r'^(\d+)\.$', cells[1])
            if not order_match:
                # Could be a continuation or field-only row
                if cells[0] and not re.match(r'^\d', cells[0]):
                    current_field = cells[0]
                # Check for continuation row (empty order, has task name in col 2)
                if not cells[1] and cells[2]:
                    if all_tasks:
                        last = all_tasks[-1]
                        last['taskName'] = clean(last['taskName'] + ' ' + cells[2])
                        if cells[3]: last['responsibleUnit'] = clean(last['responsibleUnit'] + ' ' + cells[3])
                        if cells[4]: last['coordinatingUnits'] = clean(last['coordinatingUnits'] + ' ' + cells[4])
                        if cells[5]: last['kpiCodes'] = clean(last['kpiCodes'] + ' ' + cells[5])
                        if cells[6]: last['deliverable'] = clean(last['deliverable'] + ' ' + cells[6])
                        if cells[7]: last['deadline'] = clean(last['deadline'] + ' ' + cells[7])
                continue

            order = int(order_match.group(1))

            # Field
            field = cells[0] if cells[0] else current_field
            if field:
                current_field = field

            all_tasks.append({
                'field': field,
                'order': order,
                'taskName': cells[2],
                'responsibleUnit': cells[3],
                'coordinatingUnits': cells[4],
                'kpiCodes': cells[5].replace('\u200b', '').replace('\t', ' ').strip(),
                'deliverable': cells[6],
                'deadline': cells[7],
            })

    print(f"Total raw tasks: {len(all_tasks)}")

    # Group by order resets to identify months
    month_groups = []
    current_group = []
    last_order = 0
    for task in all_tasks:
        if task['order'] <= last_order and current_group:
            month_groups.append(current_group)
            current_group = []
        current_group.append(task)
        last_order = task['order']
    if current_group:
        month_groups.append(current_group)

    print(f"Month groups: {len(month_groups)}")

    # Assign months
    result = []
    for i, group in enumerate(month_groups):
        month_key = MONTHS[i][0] if i < len(MONTHS) else f"unknown_{i}"
        for task in group:
            task_id = f"khct_{int(time.time()*1000)}_{random.randint(1000,9999)}"
            time.sleep(0.001)
            result.append({
                'id': task_id,
                'month': month_key,
                'field': task['field'],
                'order': task['order'],
                'taskName': task['taskName'],
                'responsibleUnit': task['responsibleUnit'],
                'coordinatingUnits': task['coordinatingUnits'],
                'kpiCodes': task['kpiCodes'],
                'deliverable': task['deliverable'],
                'deadline': task['deadline'],
                'status': 'active',
            })

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    with open(OUT_PATH, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"\nWritten {len(result)} tasks to {OUT_PATH}")
    from collections import Counter
    mc = Counter(t['month'] for t in result)
    for m in sorted(mc.keys(), key=lambda x: [int(p) for p in x.split('/')]):
        print(f"  {m}: {mc[m]} tasks")

if __name__ == '__main__':
    extract_khct()
